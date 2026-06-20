from docx import Document
from pathlib import Path

template_path = Path("/opt/lesson-api/templates/lesson_template.docx")

if not template_path.exists():
    print(f"模板不存在: {template_path}")
    raise SystemExit(1)

doc = Document(template_path)

print("=" * 80)
print(f"模板路径: {template_path}")
print(f"表格数量: {len(doc.tables)}")
print("=" * 80)

for table_index, table in enumerate(doc.tables):
    print(f"\n\n【表格 {table_index}】")
    print(f"行数: {len(table.rows)}")
    print(f"列数: {len(table.columns)}")
    print("-" * 80)

    for row_index, row in enumerate(table.rows):
        cell_texts = []
        for col_index, cell in enumerate(row.cells):
            text = cell.text.replace("\n", "\\n").strip()
            cell_texts.append(f"({row_index},{col_index}) {text}")

        print(" | ".join(cell_texts))
