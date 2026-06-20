from pathlib import Path

path = Path("/opt/lesson-api/main.py")
text = path.read_text(encoding="utf-8")

# 1. 确保导入分页符需要的底层 XML 工具
if "from docx.oxml import OxmlElement" not in text:
    text = text.replace(
        "from docx.oxml.ns import qn",
        "from docx.oxml.ns import qn\nfrom docx.oxml import OxmlElement"
    )

pagebreak_func = r'''
def insert_page_break_before_first_table(doc):
    """
    在第一个表格前强制插入分页符。
    用于恢复“封面”和“教案表格”之间的分页。
    """
    if not doc.tables:
        return

    first_table = doc.tables[0]
    table_element = first_table._tbl

    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r.append(br)
    p.append(r)

    table_element.addprevious(p)
'''

# 2. 插入分页函数
if "def insert_page_break_before_first_table(doc):" not in text:
    marker = "def fill_template_docx(data: LessonDocxRequest) -> str:"
    if marker not in text:
        raise SystemExit("未找到 fill_template_docx 函数，无法插入分页函数。")
    text = text.replace(marker, pagebreak_func + "\n\n" + marker)

# 3. 在 fill_cover_info 后调用分页
if "insert_page_break_before_first_table(doc)" not in text:
    text = text.replace(
        "    fill_cover_info(doc, data)\n\n    if len(doc.tables) == 0:",
        "    fill_cover_info(doc, data)\n    insert_page_break_before_first_table(doc)\n\n    if len(doc.tables) == 0:"
    )

# 4. 修改表格顶部信息：只填写左侧第一组班级和日期，不再复制到右侧
old_block = '''    # 顶部基本信息区：填写授课班级和日期
    # 时间栏只填写年月日，不填写第几节
    set_cell_text(table.cell(0, 1), value_or_default(data.class_name))
    set_cell_text(table.cell(0, 3), value_or_default(data.lesson_date))

    # 如果后续同一份教案多个班，可继续使用右侧区域
    # 当前先复制同一份班级和日期到右侧
    set_cell_text(table.cell(0, 5), value_or_default(data.class_name))
    set_cell_text(table.cell(0, 7), value_or_default(data.lesson_date))
'''

new_block = '''    # 顶部基本信息区：只填写第一组授课班级和日期
    # 同一个班级和上课时间只填写一次，不再复制到右侧区域
    set_cell_text(table.cell(0, 1), value_or_default(data.class_name))
    set_cell_text(table.cell(0, 3), value_or_default(data.lesson_date))
'''

if old_block in text:
    text = text.replace(old_block, new_block)
else:
    # 兼容早期代码块
    text = text.replace(
        '''    set_cell_text(table.cell(0, 1), value_or_default(data.class_name))
    set_cell_text(table.cell(0, 3), value_or_default(data.lesson_date))

    # 如果后续同一份教案多个班，可继续使用右侧区域
    # 当前先复制同一份信息到右侧
    set_cell_text(table.cell(0, 5), value_or_default(data.class_name))
    set_cell_text(table.cell(0, 7), value_or_default(data.lesson_date))
''',
        '''    set_cell_text(table.cell(0, 1), value_or_default(data.class_name))
    set_cell_text(table.cell(0, 3), value_or_default(data.lesson_date))
'''
    )

# 5. 更新版本号
text = text.replace('version="0.4.0"', 'version="0.5.0"')
text = text.replace('"version": "0.4.0"', '"version": "0.5.0"')
text = text.replace('version="0.3.0"', 'version="0.5.0"')
text = text.replace('"version": "0.3.0"', '"version": "0.5.0"')

path.write_text(text, encoding="utf-8")
print("更新完成：已恢复封面与表格分页，并改为只填写一次班级和日期。")
