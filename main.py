import os
from copy import deepcopy

from fastapi import FastAPI, Request
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel
from datetime import datetime
from uuid import uuid4
from pathlib import Path

from docx import Document
from docx.table import Table
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


BASE_DIR = Path(__file__).resolve().parent
OUTPUT_DIR = BASE_DIR / "generated"
TEMPLATE_DIR = BASE_DIR / "templates"
TEMPLATE_PATH = TEMPLATE_DIR / "lesson_template.docx"

OUTPUT_DIR.mkdir(exist_ok=True)
TEMPLATE_DIR.mkdir(exist_ok=True)

PUBLIC_BASE_URL = os.getenv("PUBLIC_BASE_URL", "").rstrip("/")

app = FastAPI(title="Lesson Plan API", version="0.5.0")
app.mount("/files", StaticFiles(directory=str(OUTPUT_DIR)), name="files")


class TeachingObjectives(BaseModel):
    knowledge: str | None = None
    ability: str | None = None
    quality: str | None = None


class LessonData(BaseModel):
    teaching_objectives: TeachingObjectives | None = None
    ideological_goal: str | None = None
    key_points: str | None = None
    difficult_points: str | None = None
    teaching_design: str | None = None
    reflection: str | None = None
    references: list[str] | None = None


class CourseInfo(BaseModel):
    course_name: str | None = None
    teacher_name: str | None = None
    class_name: str | None = None
    semester: str | None = None
    college: str | None = None
    teaching_group: str | None = None


class SemesterBasicInfo(BaseModel):
    college: str | None = None
    course_name: str | None = None
    teacher_name: str | None = None
    class_name: str | None = None
    semester: str | None = None
    teaching_group: str | None = None
    teaching_date: str | None = None
    period: str | None = None
    hours: str | None = None


class SemesterObjectives(BaseModel):
    knowledge_goal: str | None = None
    ability_goal: str | None = None
    quality_goal: str | None = None
    knowledge: str | None = None
    ability: str | None = None
    quality: str | None = None


class SemesterLessonContent(BaseModel):
    chapter: str | None = None
    objectives: SemesterObjectives | None = None
    ideological_goal: str | None = None
    key_points: str | None = None
    difficult_points: str | None = None
    teaching_design: str | None = None
    reflection: str | None = None
    references: list[str] | None = None


class SemesterLessonItem(BaseModel):
    lesson_no: str | None = None
    basic_info: SemesterBasicInfo | None = None
    lesson_content: SemesterLessonContent | None = None
    generation_trace: dict | None = None


class LessonDocxRequest(BaseModel):
    generation_mode: str | None = None
    course_info: CourseInfo | None = None
    lessons: list[SemesterLessonItem] | None = None
    course_name: str | None = None
    lesson_topic: str | None = None
    teacher_name: str | None = None
    class_name: str | None = None
    semester: str | None = None
    college: str | None = None
    teaching_group: str | None = None
    lesson_date: str | None = None
    lesson_time: str | None = None
    hours: str | None = None
    lesson_data: LessonData | None = None


def value_or_default(value, default=""):
    return value if value else default


def set_cell_text(cell, text: str):
    """
    清空单元格并写入文本。
    注意：模板里存在合并单元格，python-docx 仍可通过指定位置写入。
    """
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text or "")

    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(10.5)


def build_objectives_text(objectives: TeachingObjectives | None) -> str:
    objectives = objectives or TeachingObjectives()

    knowledge = value_or_default(objectives.knowledge, "掌握本次课相关基础知识、核心概念和基本原理。")
    ability = value_or_default(objectives.ability, "能够结合实际任务分析和解决本次课相关问题。")
    quality = value_or_default(objectives.quality, "培养学生严谨细致、规范操作和理论联系实际的职业素养。")

    return f"知识目标：{knowledge}\n能力目标：{ability}\n素养目标：{quality}"


def build_key_difficult_text(lesson_data: LessonData) -> str:
    key_points = value_or_default(lesson_data.key_points, "本次课相关核心概念、基本原理和应用方法。")
    difficult_points = value_or_default(lesson_data.difficult_points, "本次课抽象知识点的理解及其在实际问题中的应用。")

    return f"教学重点：{key_points}\n教学难点：{difficult_points}"


def build_references_text(references: list[str] | None) -> str:
    default_refs = [
        "[1] 康华光. 电子技术基础：模拟部分[M]. 北京: 高等教育出版社, 2021.",
        "[2] 秦曾煌. 电工学[M]. 北京: 高等教育出版社, 2018.",
        "[3] 童诗白, 华成英. 模拟电子技术基础[M]. 北京: 高等教育出版社, 2015."
    ]

    refs = references if references is not None else default_refs
    return "\n".join(refs)


def set_cover_paragraph(paragraph, label: str, value: str):
    """
    填充封面普通段落，例如：课程名称、编写教师、授课班级等。
    保留段落本身位置，重新写入“标签 + 下划线内容”。
    """
    paragraph.clear()

    label_run = paragraph.add_run(f"{label}：")
    label_run.bold = True
    label_run.font.name = "黑体"
    label_run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    label_run.font.size = Pt(16)

    value_run = paragraph.add_run(f"  {value or ''}  ")
    value_run.underline = True
    value_run.font.name = "宋体"
    value_run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    value_run.font.size = Pt(16)


def fill_cover_info(doc, data: LessonDocxRequest):
    """
    填充封面信息。
    这些内容通常不是表格，而是 Word 普通段落。
    """
    cover_map = {
        "课程名称": value_or_default(data.course_name),
        "编写教师": value_or_default(data.teacher_name),
        "授课班级": value_or_default(data.class_name),
        "授课学期": value_or_default(data.semester),
        "教研室（组）": value_or_default(getattr(data, "teaching_group", ""), ""),
        "二级学院": value_or_default(data.college)
    }

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        for label, value in cover_map.items():
            if label in text:
                set_cover_paragraph(paragraph, label, value)
                break



def insert_page_break_before_first_table(doc):
    """
    在第一个表格前强制插入分页符。
    用于恢复“封面”和“教案表格”之间的分页。
    """
    if not doc.tables:
        return

    first_table = doc.tables[0]
    table_element = first_table._tbl

    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r.append(br)
    p.append(r)

    table_element.addprevious(p)


def fill_template_docx(data: LessonDocxRequest) -> str:
    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(f"模板文件不存在: {TEMPLATE_PATH}")

    doc = Document(TEMPLATE_PATH)

    # 先填充封面信息
    fill_cover_info(doc, data)

    if len(doc.tables) == 0:
        raise ValueError("模板中没有表格，无法填充。")

    table = doc.tables[0]
    lesson_data = data.lesson_data or LessonData()

    # 顶部基本信息区：只填写第一组授课班级和日期
    # 同一个班级和上课时间只填写一次，不再复制到右侧区域
    set_cell_text(table.cell(0, 1), value_or_default(data.class_name))
    set_cell_text(table.cell(0, 3), value_or_default(data.lesson_date))

    # 主体内容区
    set_cell_text(table.cell(4, 1), value_or_default(data.lesson_topic, "未填写授课章节"))
    set_cell_text(table.cell(5, 1), build_objectives_text(lesson_data.teaching_objectives))

    ideological_goal = value_or_default(
        lesson_data.ideological_goal,
        "结合本次课程内容，引导学生树立规范意识、工程意识和责任意识，培养精益求精的职业精神。"
    )
    set_cell_text(table.cell(6, 1), ideological_goal)

    set_cell_text(table.cell(7, 1), build_key_difficult_text(lesson_data))

    teaching_design = value_or_default(
        lesson_data.teaching_design,
        "本次课围绕授课主题展开，先通过案例或问题情境导入，引导学生明确学习任务；随后讲解核心知识点，结合提问、演示、案例分析和任务训练促进学生理解；最后进行课堂总结，帮助学生梳理知识结构。"
    )
    set_cell_text(table.cell(8, 1), teaching_design)

    reflection = value_or_default(
        lesson_data.reflection,
        "本次课整体教学过程较为顺利，学生能够跟随课堂任务完成学习。后续教学中应结合学生反馈进一步优化案例设计和课堂互动，提高学生对重点难点内容的理解与应用能力。"
    )
    set_cell_text(table.cell(9, 1), reflection)

    set_cell_text(table.cell(10, 1), build_references_text(lesson_data.references))

    filename = f"template_lesson_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{uuid4().hex[:8]}.docx"
    file_path = OUTPUT_DIR / filename
    doc.save(file_path)

    return filename


@app.get("/")
def root():
    return {
        "message": "Lesson Plan API is running",
        "version": "0.5.0",
        "template_exists": TEMPLATE_PATH.exists(),
        "time": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    }


@app.get("/health")
def health():
    return {
        "status": "ok"
    }


@app.post("/generate-from-template")
def generate_from_template(data: LessonDocxRequest, request: Request):
    if data.generation_mode == "semester":
        filename = fill_semester_template_docx(data)
    else:
        filename = fill_template_docx(data)

    base_url = PUBLIC_BASE_URL or str(request.base_url).rstrip("/")
    file_url = f"{base_url}/files/{filename}"

    return {
        "success": True,
        "message": "模板教案 Word 生成成功",
        "filename": filename,
        "file_url": file_url
    }

# ===== v0.6.0 修复：封面分页 + 日期格式标准化 =====
import re as _re_for_date
from docx.enum.text import WD_BREAK


def format_chinese_date(date_text: str) -> str:
    """
    将日期统一转换为：2026 年 6 月 19 日
    兼容：
    2026/6/19
    2026-6-19
    2026.6.19
    2026年6月19日
    """
    raw = value_or_default(date_text).strip()
    if not raw:
        return ""

    match = _re_for_date.search(r"(\d{4})\D+(\d{1,2})\D+(\d{1,2})", raw)
    if match:
        year, month, day = match.groups()
        return f"{int(year)} 年 {int(month)} 月 {int(day)} 日"

    return raw


def _display_width(text: str) -> int:
    width = 0
    for char in text:
        width += 2 if ord(char) > 127 else 1
    return width


def _pad_to_display_width(text: str, target_width: int) -> str:
    current_width = _display_width(text)
    if current_width >= target_width:
        return text
    return text + (" " * (target_width - current_width))


def _center_pad_to_display_width(text: str, target_width: int) -> str:
    text = text or ""
    current_width = _display_width(text)
    if current_width >= target_width:
        return text

    padding_width = target_width - current_width
    left_padding = padding_width // 2
    right_padding = padding_width - left_padding
    return (" " * left_padding) + text + (" " * right_padding)


def _run_has_page_break(run) -> bool:
    return any(
        br.get(qn("w:type")) == "page"
        for br in run._element.xpath(".//w:br")
    )


def _paragraph_has_page_break(paragraph) -> bool:
    return any(_run_has_page_break(run) for run in paragraph.runs)


def _clear_run_text_keep_page_breaks(run):
    for node in run._element.xpath(".//w:t | .//w:tab"):
        parent = node.getparent()
        if parent is not None:
            parent.remove(node)


def _disable_proofing(run):
    rPr = run._element.get_or_add_rPr()
    no_proof = OxmlElement("w:noProof")
    rPr.append(no_proof)


def set_cover_paragraph_v2(
    paragraph,
    label: str,
    value: str,
    add_page_break: bool = False,
    max_width: int | None = None
):
    """
    填充封面普通段落，复用模板中已有的下划线空白长度。
    """
    underline_runs = [
        run for run in paragraph.runs
        if run.underline and run.text and not run.text.strip() and not _run_has_page_break(run)
    ]

    if underline_runs:
        underline_text = "".join(run.text for run in underline_runs)
        target_width = _display_width(underline_text)
        if max_width is not None:
            target_width = min(target_width, max_width)

        filled_text = _center_pad_to_display_width(value or "", target_width)

        fill_run = underline_runs[0]
        fill_run.text = filled_text
        fill_run.underline = True
        fill_run.bold = True

        for run in underline_runs[1:]:
            run.text = ""
    else:
        paragraph.clear()

        label_run = paragraph.add_run(f"{label}：")
        label_run.bold = True
        label_run.font.name = "黑体"
        label_run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        label_run.font.size = Pt(16)

        fallback_width = max_width or 32
        filled_text = _center_pad_to_display_width(value or "", fallback_width)
        value_run = paragraph.add_run(filled_text)
        value_run.underline = True
        value_run.bold = True
        value_run.font.name = "宋体"
        value_run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        value_run.font.size = Pt(16)

    if add_page_break and not _paragraph_has_page_break(paragraph):
        paragraph.add_run().add_break(WD_BREAK.PAGE)


def set_college_cover_paragraph(paragraph, value: str, add_page_break: bool = False):
    """
    单独填充“二级学院”，避免 fallback 清空段落后破坏模板分页结构。
    """
    label = "二级学院"
    label_seen = False
    page_break_run = None

    for run in paragraph.runs:
        if not label_seen:
            if label in run.text:
                run.text = f"{label}："
                run.bold = True
                label_seen = True
            continue

        if _run_has_page_break(run):
            page_break_run = page_break_run or run
            _clear_run_text_keep_page_breaks(run)
            run.underline = False
        else:
            run.text = ""

    if not label_seen:
        label_run = paragraph.add_run(f"{label}：")
        label_run.bold = True
        label_run.font.name = "黑体"
        label_run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        label_run.font.size = Pt(16)

    value_run = paragraph.add_run(f"\u3000\u3000{value or ''}\u3000\u3000")
    value_run.underline = True
    value_run.bold = True
    value_run.font.name = "黑体"
    value_run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    value_run.font.size = Pt(16)
    _disable_proofing(value_run)

    if page_break_run is not None:
        page_break_run._element.addprevious(value_run._element)
    elif add_page_break and not _paragraph_has_page_break(paragraph):
        paragraph.add_run().add_break(WD_BREAK.PAGE)


def fill_cover_info(doc, data: LessonDocxRequest):
    """
    填充封面信息。
    在“二级学院”后强制分页，使后面的教案表格从下一页开始。
    """
    cover_map = {
        "课程名称": value_or_default(data.course_name),
        "编写教师": value_or_default(data.teacher_name),
        "授课班级": value_or_default(data.class_name),
        "授课学期": value_or_default(data.semester),
        "教研室（组）": value_or_default(getattr(data, "teaching_group", "")),
        "二级学院": value_or_default(data.college)
    }

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        for label, value in cover_map.items():
            if label in text:
                if label == "二级学院":
                    set_college_cover_paragraph(paragraph, value, add_page_break=True)
                else:
                    set_cover_paragraph_v2(paragraph, label, value)
                break


def insert_page_break_before_first_table(doc):
    """
    旧版分页函数保留为空，避免重复插入分页符。
    当前分页由 fill_cover_info 在“二级学院”后完成。
    """
    return


def _create_page_break_paragraph():
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r.append(br)
    p.append(r)
    return p


def _element_has_page_break(element) -> bool:
    return any(
        br.get(qn("w:type")) == "page"
        for br in element.xpath(".//w:br")
    )


def _has_page_break_before_element(element) -> bool:
    previous = element.getprevious()
    while previous is not None:
        if previous.tag != qn("w:p"):
            return False
        if _element_has_page_break(previous):
            return True
        previous = previous.getprevious()
    return False


def ensure_page_break_before_table(table):
    if not _has_page_break_before_element(table._tbl):
        table._tbl.addprevious(_create_page_break_paragraph())


def fill_lesson_table(table, lesson: LessonDocxRequest):
    lesson_data = lesson.lesson_data or LessonData()
    formatted_date = format_chinese_date(lesson.lesson_date)

    # 顶部基本信息区：只填写第一组授课班级和日期
    # 同一个班级和上课时间只填写一次，不复制到右侧区域
    set_cell_text(table.cell(0, 1), value_or_default(lesson.class_name))
    set_cell_text(table.cell(0, 3), formatted_date)

    # 主体内容区
    set_cell_text(table.cell(4, 1), value_or_default(lesson.lesson_topic, "未填写授课章节"))
    set_cell_text(table.cell(5, 1), build_objectives_text(lesson_data.teaching_objectives))

    ideological_goal = value_or_default(
        lesson_data.ideological_goal,
        "结合本次课程内容，引导学生树立规范意识、工程意识和责任意识，培养精益求精的职业精神。"
    )
    set_cell_text(table.cell(6, 1), ideological_goal)

    set_cell_text(table.cell(7, 1), build_key_difficult_text(lesson_data))

    teaching_design = value_or_default(
        lesson_data.teaching_design,
        "本次课围绕授课主题展开，先通过案例或问题情境导入，引导学生明确学习任务；随后讲解核心知识点，结合提问、演示、案例分析和任务训练促进学生理解；最后进行课堂总结，帮助学生梳理知识结构。"
    )
    set_cell_text(table.cell(8, 1), teaching_design)

    reflection = value_or_default(
        lesson_data.reflection,
        "本次课整体教学过程较为顺利，学生能够跟随课堂任务完成学习。后续教学中应结合学生反馈进一步优化案例设计和课堂互动，提高学生对重点难点内容的理解与应用能力。"
    )
    set_cell_text(table.cell(9, 1), reflection)

    set_cell_text(table.cell(10, 1), build_references_text(lesson_data.references))


def _semester_lesson_to_docx_request(course_info: CourseInfo, lesson: SemesterLessonItem) -> LessonDocxRequest:
    basic_info = lesson.basic_info or SemesterBasicInfo()
    content = lesson.lesson_content or SemesterLessonContent()
    objectives = content.objectives or SemesterObjectives()

    teaching_objectives = TeachingObjectives(
        knowledge=objectives.knowledge_goal or objectives.knowledge,
        ability=objectives.ability_goal or objectives.ability,
        quality=objectives.quality_goal or objectives.quality
    )

    lesson_data = LessonData(
        teaching_objectives=teaching_objectives,
        ideological_goal=content.ideological_goal,
        key_points=content.key_points,
        difficult_points=content.difficult_points,
        teaching_design=content.teaching_design,
        reflection=content.reflection,
        references=content.references
    )

    return LessonDocxRequest(
        course_name=basic_info.course_name or course_info.course_name,
        lesson_topic=content.chapter,
        teacher_name=basic_info.teacher_name or course_info.teacher_name,
        class_name=basic_info.class_name or course_info.class_name,
        semester=basic_info.semester or course_info.semester,
        college=basic_info.college or course_info.college,
        teaching_group=basic_info.teaching_group or course_info.teaching_group,
        lesson_date=basic_info.teaching_date,
        lesson_time=basic_info.period,
        hours=basic_info.hours,
        lesson_data=lesson_data
    )


def fill_semester_template_docx(data: LessonDocxRequest) -> str:
    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(f"模板文件不存在: {TEMPLATE_PATH}")

    if not data.lessons:
        raise ValueError("多次课生成需要提供 lessons 数组。")

    doc = Document(TEMPLATE_PATH)

    course_info = data.course_info or CourseInfo(
        course_name=data.course_name,
        teacher_name=data.teacher_name,
        class_name=data.class_name,
        semester=data.semester,
        college=data.college,
        teaching_group=data.teaching_group
    )
    cover_data = LessonDocxRequest(
        course_name=course_info.course_name,
        teacher_name=course_info.teacher_name,
        class_name=course_info.class_name,
        semester=course_info.semester,
        college=course_info.college,
        teaching_group=course_info.teaching_group
    )

    fill_cover_info(doc, cover_data)

    if len(doc.tables) == 0:
        raise ValueError("模板中没有表格，无法填充。")

    first_table = doc.tables[0]
    blank_table_xml = deepcopy(first_table._tbl)

    ensure_page_break_before_table(first_table)
    fill_lesson_table(first_table, _semester_lesson_to_docx_request(course_info, data.lessons[0]))

    last_table_element = first_table._tbl
    for lesson in data.lessons[1:]:
        page_break = _create_page_break_paragraph()
        new_table_xml = deepcopy(blank_table_xml)

        last_table_element.addnext(page_break)
        page_break.addnext(new_table_xml)

        new_table = Table(new_table_xml, doc)
        fill_lesson_table(new_table, _semester_lesson_to_docx_request(course_info, lesson))
        last_table_element = new_table_xml

    filename = f"semester_lesson_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{uuid4().hex[:8]}.docx"
    file_path = OUTPUT_DIR / filename
    doc.save(file_path)

    return filename


def fill_template_docx(data: LessonDocxRequest) -> str:
    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(f"模板文件不存在: {TEMPLATE_PATH}")

    doc = Document(TEMPLATE_PATH)

    # 填充封面，并在封面后分页
    fill_cover_info(doc, data)

    if len(doc.tables) == 0:
        raise ValueError("模板中没有表格，无法填充。")

    table = doc.tables[0]
    fill_lesson_table(table, data)

    filename = f"template_lesson_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{uuid4().hex[:8]}.docx"
    file_path = OUTPUT_DIR / filename
    doc.save(file_path)

    return filename
# ===== v0.6.0 修复结束 =====
